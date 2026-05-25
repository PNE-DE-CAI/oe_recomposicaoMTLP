#!/usr/bin/env python3
"""
Script para extrair questões dos documentos DOCX de Língua Portuguesa e Matemática
e gerar um arquivo JSON com a estrutura das questões para integração no HTML
"""

import json
from docx import Document
from pathlib import Path
import re

def extrair_questoes_docx(caminho_arquivo):
    """
    Extrai questões de um arquivo DOCX
    Retorna lista de dicionários com as questões estruturadas
    """
    doc = Document(caminho_arquivo)
    questoes = []
    i = 0
    
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        texto = para.text.strip()
        
        # Detecta cabeçalho de questão (QUESTÃO X)
        if texto.startswith('QUESTÃO') or re.match(r'^QUESTÃO\s+\d+', texto):
            questao = {
                'numero': None,
                'titulo': '',
                'conteudo': '',
                'opcoes': [],
                'resposta': '',
                'descriptor': '',
                'justificativa': ''
            }
            
            # Extrai número da questão
            match = re.search(r'\d+', texto)
            if match:
                questao['numero'] = int(match.group())
            
            i += 1
            
            # Coleta conteúdo até encontrar as opções (a, b, c, d)
            conteudo_questao = []
            while i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                texto = para.text.strip()
                
                # Verifica se começou as opções
                if re.match(r'^[a-d]\)\s', texto):
                    break
                
                if texto and not texto.startswith('Resposta:') and not texto.startswith('Objetivo'):
                    conteudo_questao.append(texto)
                
                i += 1
            
            questao['conteudo'] = ' '.join(conteudo_questao)
            
            # Coleta as opções
            while i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                texto = para.text.strip()
                
                if re.match(r'^[a-d]\)\s', texto):
                    # Extrai a letra e o conteúdo
                    match = re.match(r'^([a-d])\)\s(.+)', texto)
                    if match:
                        letra = match.group(1)
                        conteudo = match.group(2)
                        questao['opcoes'].append(conteudo)
                    i += 1
                elif texto.startswith('Resposta:'):
                    # Extrai a resposta correta
                    match = re.search(r'Resposta:\s*([a-d])', texto)
                    if match:
                        questao['resposta'] = match.group(1)
                    i += 1
                    break
                elif texto.startswith('Objetivo'):
                    break
                else:
                    i += 1
            
            # Coleta o descriptor
            while i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                texto = para.text.strip()
                
                if texto.startswith('Objetivo'):
                    questao['descriptor'] = texto.replace('Objetivo de aprendizagem:', '').strip()
                    break
                
                i += 1
            
            if questao['numero'] and questao['opcoes']:
                questoes.append(questao)
        else:
            i += 1
    
    return questoes

def processar_todos_arquivos():
    """
    Processa todos os arquivos DOCX e retorna dicionário estruturado
    """
    diretorio = Path('/mnt/user-data/uploads')
    
    arquivos = {
        'lp': {
            '6': 'LP_6º_ano_-_Itens_OE_-_1º_bimestre.docx',
            '7': None,  # Não fornecido
            '8': None,  # Não fornecido
            '9': 'LP_9º_ano_-_Itens_OE__-_1º_bimestre.docx',
            '3': 'LP_3ª_série_-_Itens_OE_-_1º_bimestre.docx'
        },
        'mat': {
            '6': 'MAT_6º_ano_-_Itens_OE_-_1º_bimestre.docx',
            '7': None,  # Não fornecido
            '8': None,  # Não fornecido
            '9': 'MAT_9º_ano_-_Itens_OE_-_1º_bimestre.docx',
            '3': 'MAT_3ª_série_-_Itens_OE_-_1º_bimestre.docx'
        }
    }
    
    resultado = {
        'lp': {},
        'mat': {}
    }
    
    for disciplina in ['lp', 'mat']:
        for serie, arquivo in arquivos[disciplina].items():
            if arquivo:
                caminho = diretorio / arquivo
                if caminho.exists():
                    print(f"Processando {arquivo}...")
                    questoes = extrair_questoes_docx(str(caminho))
                    resultado[disciplina][serie] = {
                        'titulo': f"{'Língua Portuguesa' if disciplina == 'lp' else 'Matemática'} - {serie}º Ano" if serie != '3' else f"{'Língua Portuguesa' if disciplina == 'lp' else 'Matemática'} - 3ª Série",
                        'questoes': questoes,
                        'total': len(questoes)
                    }
                    print(f"  ✓ {len(questoes)} questões extraídas")
                else:
                    print(f"  ✗ Arquivo não encontrado: {arquivo}")
    
    return resultado

if __name__ == '__main__':
    dados = processar_todos_arquivos()
    
    # Salva em JSON
    output_path = Path('/mnt/user-data/outputs/questoes.json')
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(dados, f, ensure_ascii=False, indent=2)
    
    print(f"\n✓ Arquivo JSON salvo em: {output_path}")
    
    # Exibe resumo
    print("\n=== RESUMO ===")
    for disciplina in ['lp', 'mat']:
        print(f"\n{disciplina.upper()}:")
        for serie, dados_serie in dados[disciplina].items():
            print(f"  {serie}º Ano: {dados_serie['total']} questões")
